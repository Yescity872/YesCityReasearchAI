"""
generate_docs.py
Generates documentation.docx for the YesCity Research AI API.
Run: python generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x23, 0x23, 0x23)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x18, 0x65, 0xD6)
    return h

def add_label_value(doc, label, value, label_bold=True):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = label_bold
    run_label.font.size = Pt(11)
    run_value = p.add_run(value)
    run_value.font.size = Pt(11)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x5A)

    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F3F4')
    pPr.append(shd)

def add_param_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark blue header cell background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A73E8')
        tcPr.append(shd)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

def add_divider(doc):
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(6)
    hr.paragraph_format.space_after = Pt(6)
    run = hr.add_run('─' * 90)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

def add_method_badge(para, method):
    colors = {
        'POST': RGBColor(0xD9, 0x53, 0x1F),
        'GET':  RGBColor(0x0F, 0x9D, 0x58),
        'HEAD': RGBColor(0x74, 0x39, 0xDB),
    }
    run = para.add_run(f' {method} ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = colors.get(method.upper(), RGBColor(0x60, 0x60, 0x60))

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Cover ─────────────────────────────────────────────────────────────────────

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('YesCity Research AI')
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_para.add_run('API Route Documentation  —  For Frontend Team')
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x5F, 0x6B, 0x7C)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Generated: {datetime.date.today().strftime("%B %d, %Y")}')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

doc.add_page_break()

# ── Table of Contents ─────────────────────────────────────────────────────────

add_heading(doc, 'Table of Contents', level=1)
toc_items = [
    '1.  Overview',
    '2.  Base URL & Authentication',
    '3.  Endpoints',
    '    3.1  GET /  —  Welcome',
    '    3.2  HEAD /  —  Health Ping',
    '    3.3  GET /api/health  —  Health Check',
    '    3.4  POST /api/search  —  AI Search (POST)',
    '    3.5  GET /api/search  —  AI Search (GET)',
    '    3.6  POST /api/{user_id}/itinerary  —  Generate Itinerary (POST)',
    '    3.7  GET /api/{user_id}/itinerary  —  Generate Itinerary (GET)',
    '4.  Response Models',
    '5.  Error Handling',
    '6.  Search Category Reference',
    '7.  Itinerary Data Model',
    '8.  Quick-Start Examples',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.runs[0]
    run.font.size = Pt(11)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. Overview
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '1. Overview', level=1)
doc.add_paragraph(
    'YesCity Research AI is a FastAPI-based backend service that powers intelligent '
    'city exploration features. It exposes two main domain APIs:\n\n'
    '• Search API — classifies a natural-language query into a category (places, food, '
    'shopping, activities, accommodations, transport) and returns AI-curated recommendations.\n\n'
    '• Itinerary API — generates a personalised day-wise trip itinerary based on a user\'s '
    'saved wishlist for a given city, incorporating budget and duration extracted from the query.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. Base URL & Auth
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '2. Base URL & Authentication', level=1)
add_label_value(doc, 'Base URL (local dev)', 'http://localhost:8000')
add_label_value(doc, 'API Version', '1.0.0')
add_label_value(doc, 'Authentication', 'None (open in current version)')
add_label_value(doc, 'Content-Type', 'application/json')
add_label_value(doc, 'CORS', 'Allowed from all origins (*)')

doc.add_paragraph(
    'Note: In production, replace the wildcard CORS origin with your frontend domain.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. Endpoints
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '3. Endpoints', level=1)

# ── 3.1 GET / ─────────────────────────────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.1  Welcome', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'GET')
p.add_run('  /').font.size = Pt(12)

doc.add_paragraph('Returns a welcome message confirming the API is live.')

add_heading(doc, 'Response (200)', level=3)
add_code_block(doc, '{\n  "message": "Welcome to YesCity Search API"\n}')

# ── 3.2 HEAD / ────────────────────────────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.2  Health Ping', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'HEAD')
p.add_run('  /').font.size = Pt(12)

doc.add_paragraph(
    'Lightweight endpoint used by load balancers / uptime monitors. '
    'Returns HTTP 200 with no body.'
)

# ── 3.3 GET /api/health ───────────────────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.3  Health Check', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'GET')
p.add_run('  /api/health').font.size = Pt(12)

doc.add_paragraph('Returns service health and database connection status.')

add_heading(doc, 'Response (200)', level=3)
add_code_block(doc,
'{\n'
'  "status": "healthy",\n'
'  "database": "connected",   // "connected" | "disconnected"\n'
'  "version": "1.0.0"\n'
'}'
)

# ── 3.4 POST /api/search ──────────────────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.4  AI Search  (POST)', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'POST')
p.add_run('  /api/search').font.size = Pt(12)

doc.add_paragraph(
    'Accepts a natural-language query, classifies it into a domain category, '
    'queries MongoDB, and returns AI-generated recommendations.'
)

add_heading(doc, 'Request Body', level=3)
add_param_table(doc,
    ['Field', 'Type', 'Required', 'Description'],
    [
        ['query', 'string', '✅ Yes', 'Natural-language search query'],
    ]
)
add_code_block(doc, '{\n  "query": "best rooftop restaurants in Agra"\n}')

add_heading(doc, 'Response (200)', level=3)
add_code_block(doc,
'{\n'
'  "response": "<AI-generated recommendations string>",\n'
'  "status": "success",\n'
'  "query": "best rooftop restaurants in Agra"\n'
'}'
)

add_heading(doc, 'Error Responses', level=3)
add_param_table(doc,
    ['Status Code', 'Reason'],
    [
        ['400', 'Query is empty or whitespace'],
        ['500', 'Internal server / LLM error'],
    ]
)

# ── 3.5 GET /api/search ───────────────────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.5  AI Search  (GET)', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'GET')
p.add_run('  /api/search').font.size = Pt(12)

doc.add_paragraph(
    'Convenience GET variant of the search endpoint. Identical response schema to POST.'
)

add_heading(doc, 'Query Parameters', level=3)
add_param_table(doc,
    ['Parameter', 'Type', 'Required', 'Description'],
    [
        ['q', 'string', '✅ Yes', 'Natural-language search query (URL-encoded)'],
    ]
)

add_heading(doc, 'Example Request', level=3)
add_code_block(doc, 'GET /api/search?q=budget+hotels+in+Jaipur')

add_heading(doc, 'Response (200)', level=3)
add_code_block(doc,
'{\n'
'  "response": "<AI recommendations>",\n'
'  "status": "success",\n'
'  "query": "budget hotels in Jaipur"\n'
'}'
)

add_heading(doc, 'Error Responses', level=3)
add_param_table(doc,
    ['Status Code', 'Reason'],
    [
        ['400', "Missing or empty 'q' parameter"],
        ['500', 'Internal server / LLM error'],
    ]
)

# ── 3.6 POST /api/{user_id}/itinerary ─────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.6  Generate Itinerary  (POST)', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'POST')
p.add_run('  /api/{user_id}/itinerary').font.size = Pt(12)

doc.add_paragraph(
    'Generates a personalised, day-wise trip itinerary for the authenticated user. '
    'The backend fetches the user\'s wishlist items for the specified city from MongoDB, '
    'classifies the trip budget and duration from the free-text query, and calls an LLM '
    'to produce an optimised schedule grouped by time slots (Morning / Afternoon / Evening / Night).'
)

add_heading(doc, 'Path Parameters', level=3)
add_param_table(doc,
    ['Parameter', 'Type', 'Required', 'Description'],
    [
        ['user_id', 'string', '✅ Yes', 'MongoDB _id of the user (from your auth system)'],
    ]
)

add_heading(doc, 'Query Parameters', level=3)
add_param_table(doc,
    ['Parameter', 'Type', 'Required', 'Description'],
    [
        ['latitude',  'float',  '✅ Yes', "User's current latitude (decimal degrees, e.g. 27.1767)"],
        ['longitude', 'float',  '✅ Yes', "User's current longitude (decimal degrees, e.g. 78.0081)"],
        ['city',      'string', '✅ Yes', 'City name to generate the itinerary for (e.g. Agra)'],
        ['query',     'string', '✅ Yes', "Natural-language trip query (e.g. '3-day budget trip in Agra')"],
    ]
)

add_heading(doc, 'Example Request', level=3)
add_code_block(doc,
'POST /api/64f3a1b2c8d4e5f6a7b8c9d0/itinerary\n'
'     ?latitude=27.1767\n'
'     &longitude=78.0081\n'
'     &city=Agra\n'
'     &query=3-day+budget+trip+in+Agra'
)

add_heading(doc, 'Response (200)', level=3)
add_code_block(doc,
'{\n'
'  "user_id": "64f3a1b2c8d4e5f6a7b8c9d0",\n'
'  "city":    "Agra",\n'
'  "status":  "success",\n'
'  "data": {\n'
'    "itinerary": [\n'
'      {\n'
'        "day": 1,\n'
'        "date_label": "Day 1",\n'
'        "slots": {\n'
'          "Morning":   [{"_id": "abc123", "name": "Taj Mahal",    "category": "Place"}],\n'
'          "Afternoon": [{"_id": "def456", "name": "Agra Fort",    "category": "Place"}],\n'
'          "Evening":   [{"_id": "ghi789", "name": "Mehtab Bagh",  "category": "HiddenGem"}],\n'
'          "Night":     []\n'
'        }\n'
'      }\n'
'    ],\n'
'    "summary":  "A 3-day budget-friendly trip covering the iconic Mughal heritage sites of Agra.",\n'
'    "tips":     ["Carry cash for entry fees", "Book sunrise slots at Taj Mahal in advance"],\n'
'    "budget":   "budget",\n'
'    "duration": "3 days"\n'
'  }\n'
'}'
)

add_heading(doc, 'Error Responses', level=3)
add_param_table(doc,
    ['Status Code', 'Reason'],
    [
        ['404', 'User not found or no wishlist items for the city'],
        ['500', 'LLM or database error'],
    ]
)

# ── 3.7 GET /api/{user_id}/itinerary ──────────────────────────────────────────

add_divider(doc)
add_heading(doc, '3.7  Generate Itinerary  (GET)', level=2)
p = doc.add_paragraph()
add_method_badge(p, 'GET')
p.add_run('  /api/{user_id}/itinerary').font.size = Pt(12)

doc.add_paragraph(
    'GET convenience wrapper — identical path parameters, query parameters, '
    'and response schema to the POST variant (§ 3.6).'
)

add_heading(doc, 'Example Request', level=3)
add_code_block(doc,
'GET /api/64f3a1b2c8d4e5f6a7b8c9d0/itinerary\n'
'    ?latitude=27.1767&longitude=78.0081&city=Agra&query=weekend+trip'
)

# ─────────────────────────────────────────────────────────────────────────────
# 4. Response Models
# ─────────────────────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, '4. Response Models', level=1)

add_heading(doc, 'SearchResponse', level=2)
add_param_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['response', 'string', 'AI-generated natural-language response / recommendations'],
        ['status',   'string', '"success"'],
        ['query',    'string', 'The original query string echoed back'],
    ]
)

add_heading(doc, 'ItineraryResponse', level=2)
add_param_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['user_id', 'string', 'The user whose wishlist was used'],
        ['city',    'string', 'The target city'],
        ['status',  'string', '"success"'],
        ['data',    'object', 'The full itinerary object (see §7)'],
    ]
)

add_heading(doc, 'ErrorResponse', level=2)
add_param_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['error',  'string', 'Human-readable error description'],
        ['status', 'string', '"error"'],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. Error Handling
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '5. Error Handling', level=1)
doc.add_paragraph(
    'All endpoints follow a consistent error format. HTTP status codes used:'
)
add_param_table(doc,
    ['HTTP Code', 'Meaning', 'When it occurs'],
    [
        ['200', 'OK',                    'Request processed successfully'],
        ['400', 'Bad Request',           'Missing or invalid required parameter'],
        ['404', 'Not Found',             'User or wishlist data not found'],
        ['500', 'Internal Server Error', 'Database connection failure or LLM error'],
    ]
)

add_heading(doc, 'Error Body Example', level=3)
add_code_block(doc,
'{\n'
'  "detail": "Query cannot be empty"\n'
'}'
)
doc.add_paragraph('Note: FastAPI wraps all HTTP exceptions in a "detail" field by default.')

# ─────────────────────────────────────────────────────────────────────────────
# 6. Search Category Reference
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '6. Search Category Reference', level=1)
doc.add_paragraph(
    'The backend automatically classifies queries into one of the following categories. '
    'The frontend does NOT need to pass the category — it is inferred from the query text.'
)
add_param_table(doc,
    ['Category Key', 'Description', 'Example Query'],
    [
        ['places',         'Tourist spots, landmarks, hidden gems', '"top places to visit in Jaipur"'],
        ['foods',          'Restaurants, street food, cafes',        '"best street food in Delhi"'],
        ['shoppings',      'Markets, malls, boutiques',              '"shopping streets in Mumbai"'],
        ['activities',     'Adventures, experiences, events',        '"trekking near Manali"'],
        ['accommodations', 'Hotels, hostels, homestays',             '"budget hotels in Goa"'],
        ['transport',      'Cabs, trains, buses, logistics',         '"how to reach Shimla from Delhi"'],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 7. Itinerary Data Model
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '7. Itinerary Data Model  (data field)', level=1)
doc.add_paragraph(
    'The data field inside ItineraryResponse contains the following structure:'
)

add_param_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['itinerary', 'array',  'Ordered array of day objects'],
        ['summary',   'string', 'One-paragraph AI-generated trip overview'],
        ['tips',      'array',  'List of practical travel tips (strings)'],
        ['budget',    'string', '"budget" | "mid-range" | "luxury" | "not classified"'],
        ['duration',  'string', '"X days" | "not classified"'],
    ]
)

add_heading(doc, 'Day Object', level=2)
add_param_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['day',        'integer', 'Day number (1-indexed)'],
        ['date_label', 'string',  '"Day 1", "Day 2", etc.'],
        ['slots',      'object',  'Object with four time-slot arrays'],
    ]
)

add_heading(doc, 'Slots Object', level=2)
add_param_table(doc,
    ['Slot Key',  'Description'],
    [
        ['Morning',   '08:00 – 12:00 (sunrise views, early openings)'],
        ['Afternoon', '12:00 – 17:00 (indoor spots during hot hours)'],
        ['Evening',   '17:00 – 20:00 (sunset views, markets, light shows)'],
        ['Night',     '20:00+ (only for places explicitly open at night)'],
    ]
)

add_heading(doc, 'Place Item (inside each slot)', level=2)
add_param_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['_id',      'string', 'MongoDB ObjectId of the wishlist item'],
        ['name',     'string', 'Display name of the place'],
        ['category', 'string', 'Model type: Place | HiddenGem | NearbySpot | Shop | Activity'],
    ]
)

add_heading(doc, 'MongoDB onModel Types', level=2)
add_param_table(doc,
    ['Value', 'Meaning'],
    [
        ['Place',       'Standard tourist attraction or landmark'],
        ['HiddenGem',   'Lesser-known local favourite'],
        ['NearbySpot',  'Points of interest near the user location'],
        ['Shop',        'Shopping destination'],
        ['Activity',    'Experience or activity (tour, trek, show, etc.)'],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 8. Quick-Start Examples
# ─────────────────────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, '8. Quick-Start Examples', level=1)

add_heading(doc, 'Health Check', level=2)
add_code_block(doc,
'// JavaScript / fetch\n'
'const res = await fetch("http://localhost:8000/api/health");\n'
'const data = await res.json();\n'
'console.log(data.status); // "healthy"'
)

add_heading(doc, 'Search (POST)', level=2)
add_code_block(doc,
'const res = await fetch("http://localhost:8000/api/search", {\n'
'  method: "POST",\n'
'  headers: { "Content-Type": "application/json" },\n'
'  body: JSON.stringify({ query: "best places to visit in Agra" })\n'
'});\n'
'const data = await res.json();\n'
'console.log(data.response);'
)

add_heading(doc, 'Search (GET)', level=2)
add_code_block(doc,
"const query = encodeURIComponent('street food in Kolkata');\n"
"const res   = await fetch(`http://localhost:8000/api/search?q=${query}`);\n"
"const data  = await res.json();\n"
"console.log(data.response);"
)

add_heading(doc, 'Generate Itinerary', level=2)
add_code_block(doc,
"const userId = '64f3a1b2c8d4e5f6a7b8c9d0';\n"
"const params = new URLSearchParams({\n"
"  latitude:  '27.1767',\n"
"  longitude: '78.0081',\n"
"  city:      'Agra',\n"
"  query:     '3-day budget trip in Agra with family'\n"
"});\n\n"
"const res  = await fetch(`http://localhost:8000/api/${userId}/itinerary?${params}`);\n"
"const data = await res.json();\n\n"
"data.data.itinerary.forEach(day => {\n"
"  console.log(`Day ${day.day}:`, day.slots);\n"
"});"
)

# ─────────────────────────────────────────────────────────────────────────────
# Footer note
# ─────────────────────────────────────────────────────────────────────────────

add_divider(doc)
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(
    'YesCity Research AI  •  API v1.0.0  •  Auto-generated documentation'
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = 'documentation.docx'
doc.save(output_path)
print(f'✅  Documentation written to: {output_path}')
